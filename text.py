import math
import os.path
from nltk.corpus import stopwords

from nltk.stem import WordNetLemmatizer
from nltk import word_tokenize, sent_tokenize
import pandas as pd
from spacy import load
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
import numpy as np
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity
from sumy.parsers import plaintext
from sumy.summarizers import lsa, lex_rank, text_rank, luhn
from sumy.nlp.tokenizers import Tokenizer
from rouge_score import rouge_scorer
from scipy.signal import argrelextrema
from networkx import pagerank, from_numpy_array, Graph
from sklearn.feature_extraction.text import CountVectorizer
from pymorphy3 import MorphAnalyzer
from math import exp
from re import compile, sub
from docx import Document

SEN_PERCENT = 0.5
TOKEN = compile('\w+')
stopWords = set(stopwords.words("russian"))


class Text:
    def __init__(self, name, chapters=None):
        self.name = name
        if chapters is None:
            chapters = []
            self.chapters = chapters
        else:
            self.chapters = chapters
            timing_dict = {}
            with open(self.name + '\\timings.txt', 'r', encoding='utf-8') as file:
                while True:
                    # Читаем начальное время
                    start_time = file.readline().strip()
                    if not start_time:  # Если строка пустая, значит файл закончился
                        break
                    # Читаем текст
                    text = file.readline().strip()
                    # Читаем конечное время
                    end_time = file.readline().strip()
                    # Добавляем данные в словарь
                    timing_dict[start_time] = (text, end_time)
                    # Пропускаем пустую строку
                    file.readline()
            chapters = self.find_chapter_for_text(timing_dict)
            self.chapters_text = self.find_chapter_for_text(timing_dict)


        recognized_text = ''
        with open(name + '\\' + name + '.txt', encoding='utf-8') as file_text:
            chunk = file_text.read(10000)
            while chunk:
                recognized_text = recognized_text + chunk
                chunk = file_text.read(10000)
        self.text = recognized_text
        self.text_no_punkt = " ".join(TOKEN.findall(self.text))
        # self.data = {"Исходный текст": [recognized_text], "Без пунктуации": [text_no_punkt]}

        self.tokenize()
        self.paragraphs, self.sents_per_para = self.split_paragraphs()
        self.morph = MorphAnalyzer()
        self.model = load('ru_core_news_md', exclude=['parser', 'attribute_ruler', 'morphologizer'])
        self.sentences_count = math.ceil(SEN_PERCENT * len(self.sentences))
        self.ents, self.freqTable = self.preproccess()


    def tokenize(self):
        sentences = sent_tokenize(self.text)
        self.sentences = [sent[:len(sent) - 1] for sent in sentences]
        tokens = word_tokenize(self.text_no_punkt.lower(), language='russian')
        tokens = [token for token in tokens if token not in stopWords]
        self.tokens = list(dict.fromkeys(tokens))
        print("Количество предложений исходного текста:", len(self.sentences))

    def split_paragraphs(self):
        print('Разделение текста на параграфы')
        model = SentenceTransformer('DiTy/bi-encoder-russian-msmarco')
        sentence_length = [len(each) for each in self.sentences]
        # long = np.mean(sentence_length) + np.std(sentence_length) * 2
        # text = ''
        # for each in self.sentences:
        #     if len(each) > long:
        #         # let's replace all the commas with dots
        #         each = each.replace(',', '.')
        #     text += f'{each}. '
        # sentences = text.split('. ')
        # Embed sentences
        embeddings = model.encode(self.sentences)
        print(embeddings.shape)

        # Normalize the embeddings
        norms = np.linalg.norm(embeddings, axis=1, keepdims=True)
        embeddings = embeddings / norms
        # Create similarities matrix
        similarities = cosine_similarity(embeddings)
        activated_similarities = activate_similarities(similarities, p_size=10)

        minmimas = argrelextrema(activated_similarities, np.less, order=2)
        # Create empty string
        split_points = [each for each in minmimas[0]]
        text = ''
        paragraphs = []
        para = ''
        count = 0
        sent_para = []
        for num, sen in enumerate(self.sentences):
            if num in split_points:
                count += 1
                paragraphs.append(para)
                para = f'{sen}. '
                text += f'\n {sen}. '
            else:
                text += f'{sen}. '
                para += f'{sen}. '
            sent_para.append((sen, count))

        print('paragraphs', len(paragraphs))
        self.export_to_doc('Исходный текст', paragraphs)
        # print('\n'.join(paragraphs))
        with open(self.name + "\\paragraphs.txt", 'w', encoding="utf-8") as f:
            f.write(text)
        return paragraphs, sent_para

    def morphy_lemmatize(self, text):
        words = text.split()  # разбиваем текст на слова
        if len(words) > 1:
            res = list()
            for word in words:
                p = self.morph.parse(word.lower())[0]
                res.append(p.normal_form)
                return res
        else:
            p = self.morph.parse(words[0].lower())[0]
            return p.normal_form

    def sent_para_summary(self, filename, sentences):
        if not sentences:
            return
        parag = []
        p = ''
        with open(self.name + "\\" + filename + '.txt', 'w', encoding='utf-8') as f:
            for para in self.paragraphs:
                sen_para = sent_tokenize(para, language='russian')
                for sentence in sen_para:
                    if sentence[:-1] in sentences:
                        f.write(sentence + " ")
                        p += sentence + " "
                f.write('\n\n')
                parag.append(p)
                p = ''
        self.export_to_doc(filename, parag)

    def preproccess(self):
        if not self.tokens or not self.sentences:
            return
        doc = self.model(self.text_no_punkt.lower())
        ents = set(str(ent).lower() for ent in doc.ents)
        lemmas = []
        freqTable = dict()
        for word in self.tokens:
            if word in stopWords:
                continue
            word = self.morphy_lemmatize(word).lower()
            if word in freqTable:
                freqTable[word] += 1
            else:
                freqTable[word] = 1
                lemmas.append(word)
        # self.data['Лемматизированный текст pymorphy'] = [lemmas]
        return ents, freqTable

    def sent_summary(self):
        sentenceValue = dict()
        allFreq = 0
        try:
            for sentence in self.sentences:
                tokens = word_tokenize(" ".join(TOKEN.findall(sentence)), language='russian')
                for word in tokens:
                    word = self.morphy_lemmatize(word).lower()
                    if word in self.freqTable:
                        if sentence in sentenceValue:
                            sentenceValue[sentence] += self.freqTable[word]
                        else:
                            sentenceValue[sentence] = self.freqTable[word]
                allFreq += sentenceValue[sentence]
        except Exception as e:
            print(e)
            print(repr(e))

        average = int(allFreq / len(sentenceValue))

        # Storing sentences into our summary.
        try:
            with open(self.name + "\\" + "sent_summary.txt", 'w', encoding='utf-8') as f:
                for sentence in self.sentences:
                    if sentence in sentenceValue:
                        if (sentenceValue[sentence] > ((1 + SEN_PERCENT) * average)) or (self.contains_word(sentence)):
                            f.write(sentence.strip() + ". ")
        except Exception as e:
            print(e)
            print(repr(e))
            print(type(e).__name__)

    # def text_rank(self):
    #     sentence_paragraph_map = []
    #     for para_idx, paragraph in enumerate(self.paragraphs):
    #         para_sentences = sent_tokenize(paragraph)
    #         sentence_paragraph_map.extend([para_idx] * len(para_sentences))
    #     num_sentences = math.ceil(len(self.sentences) * SEN_PERCENT)
    #     vectorizer = CountVectorizer().fit_transform(self.sentences)
    #     vectors = vectorizer.toarray()
    #     cosine_matrix = cosine_similarity(vectors)
    #
    #     graph = from_numpy_array(cosine_matrix)
    #     scores = pagerank(graph)
    #     ranked_sentences = sorted(((score, idx) for idx, score in scores.items()), reverse=True)
    #
    #     top_sentence_indices = [idx[1] for idx in ranked_sentences[:num_sentences]]
    #     top_sentence_indices.sort()
    #
    #     summary_paragraphs = [[] for _ in range(len(self.paragraphs))]
    #
    #     for idx in top_sentence_indices:
    #         para_idx = sentence_paragraph_map[idx]
    #         summary_paragraphs[para_idx].append(self.sentences[idx])
    #
    #     summary = '\n\n'.join([' '.join(paragraph) for paragraph in summary_paragraphs if paragraph])
    #     self.export_to_doc(summary_paragraphs)
    #     with open(self.name + '\\text_rank.txt', 'w', encoding='utf-8') as f:
    #         f.write(summary)
    #     # self.data['text_rank'] = [summary]
    #     return summary

    def sumy_sum(self):
        # , {k: v for k, v in freqTable.items() if v != 1}
        lsa_sum = lsa.LsaSummarizer()
        lex_sum = lex_rank.LexRankSummarizer()
        text_rank_sum = text_rank.TextRankSummarizer()
        luhn_sum = luhn.LuhnSummarizer()
        text = ''
        parser = plaintext.PlaintextParser.from_file(self.name + "\\" + self.name + ".txt", Tokenizer('russian'))

        for sen in lsa_sum(parser.document, self.sentences_count):
            text += str(sen) + ' '
        # self.data['lsa'] = [text]
        self.sent_para_summary("lsa", text)

        text = ''
        for sen in lex_sum(parser.document, self.sentences_count):
            text += str(sen) + ' '
        # self.data['lex_rank'] = [text]
        self.sent_para_summary("lex_rank", text)
        text = ''

        for sen in text_rank_sum(parser.document, self.sentences_count):
            text += str(sen) + ' '
        # self.data['text_rank_sumy'] = [text]
        self.sent_para_summary("text_rank_sumy", text)
        text = ''

        for sen in luhn_sum(parser.document, self.sentences_count):
            text += str(sen) + ' '
        # self.data['luhn'] = [text]
        self.sent_para_summary("luhn", text)

    def contains_word(self, sentence):
        sentence_words = set(sentence.lower().split())
        words_set = set(str(word).lower() for word in self.ents)
        return not sentence_words.isdisjoint(words_set)

    def compare_sum(self):
        scorer = rouge_scorer.RougeScorer(['rouge2', 'rougeL'], use_stemmer=True)
        sums = ['lsa', 'luhn', 'sent_summary', 'text_rank_sumy', 'lex_rank']
        all_scores = [[] for _ in range(len(sums))]
        for i in range(len(sums)):
            for j in range(len(sums)):
                with open(self.name + '\\' + sums[i] + '.txt', 'r', encoding='utf-8') as t:
                    target = t.read()
                with open(self.name + '\\' + sums[j] + '.txt', 'r', encoding='utf-8') as p:
                    prediction = p.read()

                scores = scorer.score(target, prediction)
                elem = [(scores['rouge2'][0], scores['rouge2'][1], scores['rouge2'][2]),
                        ((scores['rougeL'][0], scores['rougeL'][1], scores['rougeL'][2]))]
                all_scores[i].append(elem)
        df = pd.DataFrame(all_scores)
        df.to_excel(excel_writer=self.name + '\\scores.xlsx')
        self.scores = all_scores
        print(all_scores)

    def add_data_export(self, dataset):
        print()
        if (os.path.isfile(dataset)):
            df = pd.read_csv(dataset, index_col=0)
            if self.data['Исходный текст'] not in df['Исходный текст'].values:
                df_copy = pd.DataFrame(self.data, index=[1])
                df = pd.concat([df_copy, df], ignore_index=True)
                if 'Unnamed: 0' in df.columns:
                    df.drop(columns=['Unnamed: 0'], inplace=True)
        else:
            df = pd.DataFrame(self.data, index=[0])
            df.to_csv(dataset)

    def export_to_doc(self, name, paragraphs):
        doc = Document()
        is_para_in_chapter = False
        head = doc.add_heading(self.name[0].upper() + self.name[1:], level=1)
        head.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if len(self.chapters) > 1:
            timing_dict = {}
            # Добавление заголовка и абзаца
            summary_chapters = {}
            for chapter, content in self.chapters_text.items():
                for paragraph in paragraphs:
                    paragraph = sent_tokenize(paragraph, language='russian')
                    for sentence in paragraph:
                        if sentence in content:
                            is_para_in_chapter = True
                            if chapter in summary_chapters:
                                summary_chapters[chapter] = summary_chapters[chapter] + " " + sentence
                            else:
                                summary_chapters[chapter] = sentence
                    if is_para_in_chapter:
                        summary_chapters[chapter] = summary_chapters[chapter] + "\n"
                    is_para_in_chapter = False
            for chapter, content in summary_chapters.items():
                head = doc.add_heading(chapter[0].upper() + chapter[1:], level= 2 )
                font_head = head.style.font
                font_head.name = 'Times New Roman'
                font_head.size = Pt(12)  # Устанавливаем размер шрифта
                font_head.color.rgb =  RGBColor(0, 0, 0)
                para = doc.add_paragraph(content)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                font = para.style.font
                font.name = 'Times New Roman'  # Устанавливаем шрифт
                font.size = Pt(12)
        else:
            for para in paragraphs:
                para = doc.add_paragraph(para)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                font = para.style.font
                font.name = 'Times New Roman'  # Устанавливаем шрифт
                font.size = Pt(12)
        doc.save(self.name + '\\' + name + '.docx')

    def find_chapter_for_text(self, timing_dict):
        chapters_sents = {}
        for start_time, (text, end_time) in timing_dict.items():
            start_time = float(int(start_time) / 100)
            end_time = int(int(end_time) / 100)
            for index, chapter in enumerate(self.chapters):
                if end_time <= chapter['end_time'] and start_time >= chapter['start_time']:
                    title = chapter['title']
                    if title in chapters_sents:
                        chapters_sents[title] = chapters_sents[title] + ' ' + text
                    else:
                        chapters_sents[title] = text
                    break
        return chapters_sents


def read_subs(folder):
    name = folder + '\\' + 'sub.srt.ru.vtt'
    recognized_text = ''
    with open(name) as f:
        chunk = f.read(10000)
        while chunk:
            recognized_text = recognized_text + chunk
            chunk = f.read(10000)
    new_file = folder + '\\' + 'new_sub.srt.ru.vtt'
    new_text = sub(r'(\d{2}:\d{2}:\d{2}\.\d{3})|(<c>)|(</c>)|(align:start position:0%)|(-->)|(\s{2})', '',
                   recognized_text)
    new_text = sub(r'(\s{3})|(<>)', '', new_text)
    new_text = new_text.replace('WEBVTT '
                                'Kind: captions'
                                'Language: ru', "")

    with open(new_file, 'w') as new_file:
        new_file.write(new_text)
        print('Сохранено')


def build_similarity_matrix(sentences):
    vectorizer = CountVectorizer().fit_transform(sentences)
    vectors = vectorizer.toarray()
    similarity_matrix = cosine_similarity(vectors)
    return similarity_matrix


def pagerank_sentences(sentences, similarity_matrix):
    nx_graph = from_numpy_array(similarity_matrix)
    scores = pagerank(nx_graph)
    ranked_sentences = sorted(((scores[i], s) for i, s in enumerate(sentences)), reverse=True)
    return ranked_sentences


def rev_sigmoid(x: float) -> float:
    return (1 / (1 + exp(0.5 * x)))


def activate_similarities(similarities: np.array, p_size=10) -> np.array:
    """ Function returns list of weighted sums of activated sentence similarities
    Args:
        similarities (numpy array): it should square matrix where each sentence corresponds to another with cosine similarity
        p_size (int): number of sentences are used to calculate weighted sum
    Returns:
        list: list of weighted sums
    """
    # To create weights for sigmoid function we first have to create space. P_size will determine number of sentences used and the size of weights vector.
    x = np.linspace(-10, 10, p_size)
    # Then we need to apply activation function to the created space
    y = np.vectorize(rev_sigmoid)
    # Because we only apply activation to p_size number of sentences we have to add zeros to neglect the effect of every additional sentence and to match the length ofvector we will multiply
    activation_weights = np.pad(y(x), (0, similarities.shape[0] - p_size))
    ### 1. Take each diagonal to the right of the main diagonal
    diagonals = [similarities.diagonal(each) for each in range(0, similarities.shape[0])]
    ### 2. Pad each diagonal by zeros at the end. Because each diagonal is different length we should pad it with zeros at the end
    diagonals = [np.pad(each, (0, similarities.shape[0] - len(each))) for each in diagonals]
    ### 3. Stack those diagonals into new matrix
    diagonals = np.stack(diagonals)
    ### 4. Apply activation weights to each row. Multiply similarities with our activation.
    diagonals = diagonals * activation_weights.reshape(-1, 1)
    ### 5. Calculate the weighted sum of activated similarities
    activated_similarities = np.sum(diagonals, axis=0)
    return activated_similarities


def sentence_similarity(sent1, sent2):
    words1 = set(sent1)
    words2 = set(sent2)

    common_words = words1.intersection(words2)
    if not common_words:
        return 0
    return len(common_words) / (len(words1) + len(words2))


def build_graph(sentences):
    graph = Graph()
    for i, sent1 in enumerate(sentences):
        for j, sent2 in enumerate(sentences):
            if i != j:
                similarity = sentence_similarity(sent1, sent2)
                if similarity > 0:
                    graph.add_edge(i, j, weight=similarity)
    return graph


def rank_sentences(graph):
    scores = pagerank(graph, weight='weight')
    ranked_sentences = sorted(((score, idx) for idx, score in scores.items()), reverse=True)
    return ranked_sentences


def tokenize_and_lemmatize(sentence):
    lemmatizer = WordNetLemmatizer()
    tokens = word_tokenize(sentence)
    lemmas = [lemmatizer.lemmatize(token.lower()) for token in tokens if token.isalnum() and
              token.lower() not in stopWords]
    return lemmas
